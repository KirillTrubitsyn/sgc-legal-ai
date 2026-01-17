"""
DOCX generator service for exporting AI responses
Стиль: Правовое заключение (корпоративный юридический документ)
"""
import io
import re
import os
from datetime import datetime
from typing import Optional, List, Tuple
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# Путь к логотипу
LOGO_PATH = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), 'static', 'logo.png')


def create_response_docx(
    question: str,
    answer: str,
    model: Optional[str] = None,
    created_at: Optional[datetime] = None
) -> bytes:
    """
    Create a DOCX document in analytical brief style
    """
    doc = Document()

    # Set up page margins (left 3cm, others 2cm)
    for section in doc.sections:
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

    # Set up styles
    _setup_styles(doc)

    # Header with logo
    if os.path.exists(LOGO_PATH):
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_para.add_run()
        run.add_picture(LOGO_PATH, width=Inches(0.6))
        header_para.paragraph_format.space_after = Pt(4)

    # Main title - "ПРАВОВОЕ ЗАКЛЮЧЕНИЕ"
    title = doc.add_paragraph()
    title_run = title.add_run("ПРАВОВОЕ ЗАКЛЮЧЕНИЕ")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.name = 'Times New Roman'
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.space_before = Pt(0)

    # Extract sources from text and clean answer
    clean_answer, sources = _extract_sources(answer)

    # Main content - parsed answer
    _add_formatted_text(doc, clean_answer)

    # Sources section if any
    if sources:
        _add_sources_section(doc, sources)

    # Footer
    _add_footer(doc, model, created_at)

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()


def _extract_sources(text: str) -> Tuple[str, List[str]]:
    """
    Extract [N] references from text and return clean text + list of sources.
    """
    sources = []

    # Find all [N] or [N][M] patterns and surrounding context
    # Pattern to find source indicators like [1], [2][3], etc.
    source_pattern = r'\[(\d+)\]'

    # Find unique source numbers
    found_numbers = set(re.findall(source_pattern, text))

    # Remove [N] patterns from text
    clean_text = re.sub(r'\s*\[\d+\](?:\[\d+\])*', '', text)

    # Clean up multiple spaces
    clean_text = re.sub(r'  +', ' ', clean_text)

    # Generate placeholder sources (since actual URLs aren't in the text)
    for num in sorted(found_numbers, key=int):
        sources.append(f"[{num}] Источник из результатов поиска Perplexity")

    return clean_text, sources


def _add_sources_section(doc: Document, sources: List[str]):
    """Add sources section at the end"""
    # Section header
    doc.add_paragraph()
    header = doc.add_paragraph()
    header_run = header.add_run("Источники")
    header_run.bold = True
    header_run.font.size = Pt(11)
    header_run.font.name = 'Times New Roman'
    header.paragraph_format.space_before = Pt(8)
    header.paragraph_format.space_after = Pt(4)

    # List sources
    for source in sources:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(source)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.italic = True


def _add_footer(doc: Document, model: Optional[str], created_at: Optional[datetime]):
    """Add footer with date in document and disclaimer in page footer"""
    doc.add_paragraph()

    # Separator
    separator = doc.add_paragraph()
    sep_run = separator.add_run("─" * 50)
    sep_run.font.size = Pt(8)
    separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
    separator.paragraph_format.space_after = Pt(4)
    separator.paragraph_format.space_before = Pt(4)

    # Date
    if created_at:
        date_str = created_at.strftime("%d.%m.%Y")
    else:
        date_str = datetime.now().strftime("%d.%m.%Y")

    meta_para = doc.add_paragraph()
    meta_run = meta_para.add_run(f"Дата: {date_str}")
    meta_run.font.size = Pt(9)
    meta_run.font.name = 'Times New Roman'
    meta_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    meta_para.paragraph_format.space_after = Pt(2)

    # Disclaimer in page footer (колонтитул)
    disclaimer_text = "Документ подготовлен SGC Legal AI"
    if model:
        disclaimer_text += f" | {_format_model_name(model)}"

    # Add to actual page footer
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.clear()
        footer_run = footer_para.add_run(disclaimer_text)
        footer_run.font.size = Pt(8)
        footer_run.font.name = 'Times New Roman'
        footer_run.italic = True
        footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _format_model_name(model: str) -> str:
    """Format model ID to readable name"""
    model_names = {
        # Single mode - показываем режим вместо модели
        "google/gemini-3-flash-preview": "Быстрый",
        "google/gemini-3-pro-preview": "Думающий",
        # Consilium
        "consilium": "Консилиум",
        # Fallback для других моделей
        "anthropic/claude-opus-4.5": "Claude Opus 4.5",
        "anthropic/claude-sonnet-4": "Claude Sonnet 4",
        "openai/gpt-5.2": "GPT 5.2",
        "openai/gpt-4o": "GPT 4o",
        "perplexity/sonar-pro-search": "Perplexity Sonar",
    }
    return model_names.get(model, model.split("/")[-1])


def _setup_styles(doc: Document):
    """Set up document styles"""
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(4)


def _clean_text_for_docx(text: str) -> str:
    """Remove duplicate headers and clean up text before parsing"""
    # Remove "ПРАВОВОЕ ЗАКЛЮЧЕНИЕ" header anywhere in text (on its own line)
    text = re.sub(r'^[\s]*ПРАВОВОЕ ЗАКЛЮЧЕНИЕ[^\n]*\n?', '', text, flags=re.MULTILINE | re.IGNORECASE)

    # Remove "АНАЛИТИЧЕСКАЯ СПРАВКА" - either as header on its own line, or as inline prefix
    text = re.sub(r'^[\s]*АНАЛИТИЧЕСКАЯ СПРАВКА[:\s]*', '', text, flags=re.MULTILINE | re.IGNORECASE)

    # Remove "Председатель юридического консилиума" and similar signatures
    text = re.sub(r'^Председатель\s+(юридического\s+)?консилиума.*$', '', text, flags=re.MULTILINE | re.IGNORECASE)

    # Remove "Дата составления заключения" lines
    text = re.sub(r'^Дата составления заключения.*$', '', text, flags=re.MULTILINE | re.IGNORECASE)

    # Remove --- separators
    text = re.sub(r'^---+\s*$', '', text, flags=re.MULTILINE)

    # Split inline numbered lists onto separate lines (e.g., "1. Text 2. Text" -> "1. Text\n2. Text")
    # This handles cases where conclusions are on one line
    # Match any non-space character followed by space and a number with period
    text = re.sub(r'([.,:;!?])\s+(\d+)\.\s+', r'\1\n\n\2. ', text)

    # Remove markdown bold markers **text** -> text (keep only section headers bold)
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)

    # Remove trailing empty lines
    text = re.sub(r'\n{3,}', '\n\n', text)

    return text.strip()


def _add_formatted_text(doc: Document, text: str):
    """Add text with formatting"""
    # Clean text first
    text = _clean_text_for_docx(text)

    lines = text.split('\n')
    current_para = None
    in_list = False

    for line in lines:
        stripped = line.strip()

        if not stripped:
            current_para = None
            in_list = False
            continue

        # Section headers with letters (A., B., C., D. TITLE)
        letter_section_match = re.match(r'^([A-ZА-Я])\.\s+(.*)$', stripped)
        if letter_section_match and len(stripped) < 100:
            letter, title = letter_section_match.groups()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            current_para = None
            in_list = False
            continue

        # Section headers (N. Title or N.N. Title)
        section_match = re.match(r'^(\d+(?:\.\d+)?)\.\s+([А-ЯA-Z].*)$', stripped)
        if section_match and len(stripped) < 100:
            num, title = section_match.groups()
            p = doc.add_paragraph()
            # Меньше отступы для подразделов
            if '.' in num:
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(11)
            else:
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            current_para = None
            in_list = False
            continue

        # Subsection headers (short lines starting with capital letter, no period at end)
        # Examples: "Применимые нормы права", "Анализ правоотношений"
        if (len(stripped) < 60 and
            stripped[0].isupper() and
            not stripped.endswith('.') and
            not stripped.endswith(':') and
            not re.match(r'^\d', stripped) and
            not stripped.startswith(('-', '•', '*'))):
            # Check if it looks like a heading (mostly letters, no long sentences)
            words = stripped.split()
            if len(words) <= 6 and all(not w.endswith(',') for w in words[:-1]):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'
                current_para = None
                in_list = False
                continue

        # Markdown headers (####, ###, ##, #)
        if stripped.startswith('#### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(stripped[5:])
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
            current_para = None
            continue
        elif stripped.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(stripped[4:])
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
            current_para = None
            continue
        elif stripped.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(stripped[3:])
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
            current_para = None
            continue
        elif stripped.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(stripped[2:])
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            current_para = None
            continue

        # Bullet lists
        if stripped.startswith(('- ', '• ', '* ')):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.7)
            p.paragraph_format.first_line_indent = Cm(-0.4)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            bullet_run = p.add_run("• ")
            bullet_run.font.name = 'Times New Roman'
            bullet_run.font.size = Pt(11)
            _add_inline_formatting(p, stripped[2:])
            current_para = None
            in_list = True
            continue

        # Numbered items in inline lists (e.g., "1. Объекты... 2. Объекты...")
        # Check if line contains multiple numbered items
        if re.search(r'\d+\.\s+\*\*[^*]+\*\*:', stripped):
            _add_numbered_inline_list(doc, stripped)
            current_para = None
            in_list = False
            continue

        # Regular paragraph
        if current_para is None or in_list:
            current_para = doc.add_paragraph()
            current_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            current_para.paragraph_format.first_line_indent = Cm(1.25)
            current_para.paragraph_format.space_after = Pt(6)
            current_para.paragraph_format.space_before = Pt(0)
            in_list = False
        else:
            current_para.add_run(' ')

        _add_plain_text(current_para, stripped)


def _add_numbered_inline_list(doc: Document, text: str):
    """Split inline numbered list into separate lines"""
    # Pattern to find numbered items like "1. **Something**: text"
    pattern = r'(\d+\.\s+)'
    parts = re.split(pattern, text)

    current_text = ""
    for i, part in enumerate(parts):
        if re.match(r'^\d+\.\s+$', part):
            # This is a number prefix, save it
            if current_text.strip():
                # Add previous item
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                _add_inline_formatting(p, current_text.strip())
            current_text = part
        else:
            current_text += part

    # Add last item
    if current_text.strip():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _add_inline_formatting(p, current_text.strip())


def _add_plain_text(paragraph, text: str):
    """Add text with markdown bold/italic formatting converted to DOCX"""
    _add_inline_formatting(paragraph, text)


def _add_inline_formatting(paragraph, text: str):
    """Add plain text without markdown formatting (bold only for section headers)"""
    # Remove any remaining markdown bold/italic markers
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'\*([^*]+)\*', r'\1', text)

    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
