"""
File processing service for multimodal input
"""
import os
import tempfile
import base64
from typing import Tuple
import docx
import pdfplumber
import fitz  # PyMuPDF
from openpyxl import load_workbook
from io import BytesIO
import requests
from app.config import settings


def detect_file_type(filename: str) -> str:
    """Определить тип файла по расширению"""
    ext = filename.lower().split('.')[-1]

    if ext in ['docx', 'doc']:
        return 'document'
    elif ext == 'pdf':
        return 'pdf'
    elif ext in ['xlsx', 'xls', 'xlsm']:
        return 'spreadsheet'
    elif ext in ['txt', 'md', 'markdown']:
        return 'text'
    elif ext in ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff']:
        return 'image'
    elif ext in ['mp3', 'wav', 'ogg', 'm4a', 'webm']:
        return 'audio'
    else:
        return 'unknown'


def get_audio_mime_type(filename: str) -> str:
    """Получить MIME тип для аудио файла"""
    ext = filename.lower().split('.')[-1]
    mime_types = {
        'mp3': 'audio/mpeg',
        'wav': 'audio/wav',
        'ogg': 'audio/ogg',
        'm4a': 'audio/mp4',
        'webm': 'audio/webm'
    }
    return mime_types.get(ext, 'audio/mpeg')


def get_image_mime_type(filename: str) -> str:
    """Получить MIME тип для изображения"""
    ext = filename.lower().split('.')[-1]
    mime_types = {
        'jpg': 'image/jpeg',
        'jpeg': 'image/jpeg',
        'png': 'image/png',
        'gif': 'image/gif',
        'bmp': 'image/bmp',
        'tiff': 'image/tiff',
        'webp': 'image/webp'
    }
    return mime_types.get(ext, 'image/png')


async def process_file(file_content: bytes, filename: str) -> Tuple[str, str]:
    """
    Обработать файл и извлечь текст
    Returns: (extracted_text, file_type)
    """
    file_type = detect_file_type(filename)

    if file_type == 'document':
        text = extract_docx(file_content)
    elif file_type == 'pdf':
        text = await extract_pdf(file_content)
    elif file_type == 'spreadsheet':
        text = extract_excel(file_content)
    elif file_type == 'text':
        text = file_content.decode('utf-8', errors='ignore')
    elif file_type == 'image':
        text = await extract_image_gemini(file_content, filename)
    elif file_type == 'audio':
        text = await transcribe_audio_gemini(file_content, filename)
    else:
        raise ValueError(f"Неподдерживаемый тип файла: {filename}")

    return text, file_type


def extract_docx(content: bytes) -> str:
    """Извлечь текст из DOCX"""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    try:
        doc = docx.Document(tmp_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Также извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    paragraphs.append(row_text)

        return '\n\n'.join(paragraphs)
    finally:
        os.unlink(tmp_path)


def extract_excel(content: bytes) -> str:
    """
    Извлечь данные из Excel файла (xlsx, xls, xlsm)
    Конвертирует каждый лист в markdown таблицу
    """
    wb = load_workbook(filename=BytesIO(content), read_only=True, data_only=True)
    result_parts = []

    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        rows = list(sheet.iter_rows(values_only=True))

        if not rows:
            continue

        # Фильтруем полностью пустые строки
        rows = [row for row in rows if any(cell is not None for cell in row)]

        if not rows:
            continue

        # Определяем максимальное количество колонок
        max_cols = max(len(row) for row in rows)

        # Нормализуем строки до одинаковой длины
        normalized_rows = []
        for row in rows:
            normalized_row = list(row) + [None] * (max_cols - len(row))
            normalized_rows.append(normalized_row)

        # Формируем markdown таблицу
        md_lines = []
        md_lines.append(f"## Лист: {sheet_name}\n")

        # Заголовок таблицы (первая строка)
        header = normalized_rows[0]
        header_cells = [str(cell) if cell is not None else '' for cell in header]
        md_lines.append("| " + " | ".join(header_cells) + " |")

        # Разделитель
        md_lines.append("| " + " | ".join(["---"] * max_cols) + " |")

        # Данные (остальные строки)
        for row in normalized_rows[1:]:
            row_cells = [str(cell) if cell is not None else '' for cell in row]
            md_lines.append("| " + " | ".join(row_cells) + " |")

        result_parts.append("\n".join(md_lines))

    wb.close()

    if not result_parts:
        raise ValueError("Excel файл пустой или не содержит данных")

    return "\n\n".join(result_parts)


async def extract_pdf(content: bytes) -> str:
    """
    Извлечь текст из PDF
    Сначала пробует pdfplumber, если текст не извлёкся - использует OCR через Gemini
    """
    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    try:
        # Попробовать извлечь текст напрямую
        text_parts = []
        with pdfplumber.open(tmp_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

                # Извлекаем таблицы
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        row_text = ' | '.join([str(cell) if cell else '' for cell in row])
                        if row_text.strip():
                            text_parts.append(row_text)

        extracted_text = '\n\n'.join(text_parts)

        # Если текст извлёкся - возвращаем его
        if extracted_text and len(extracted_text.strip()) > 50:
            return extracted_text

        # Если текст не извлёкся - это сканированный PDF, используем OCR
        return await extract_pdf_ocr(tmp_path)

    finally:
        os.unlink(tmp_path)


async def extract_pdf_ocr(pdf_path: str) -> str:
    """
    OCR для сканированного PDF через Gemini
    Конвертирует страницы в изображения и отправляет на распознавание
    """
    text_parts = []

    # Открываем PDF с помощью PyMuPDF
    doc = fitz.open(pdf_path)

    for page_num in range(len(doc)):
        page = doc[page_num]

        # Рендерим страницу в изображение (300 DPI для хорошего качества OCR)
        mat = fitz.Matrix(2.0, 2.0)  # 2x zoom ≈ 144 DPI (достаточно для OCR)
        pix = page.get_pixmap(matrix=mat)
        img_bytes = pix.tobytes("png")

        # Отправляем на OCR
        page_text = await ocr_image_gemini(img_bytes)
        if page_text and page_text.strip() and page_text != "Текст не обнаружен":
            text_parts.append(f"--- Страница {page_num + 1} ---\n{page_text}")

    doc.close()

    if not text_parts:
        raise ValueError("Не удалось извлечь текст из PDF (OCR не обнаружил текст)")

    return '\n\n'.join(text_parts)


async def ocr_image_gemini(image_bytes: bytes) -> str:
    """OCR для одного изображения через Gemini"""
    image_base64 = base64.b64encode(image_bytes).decode('utf-8')

    response = requests.post(
        "https://openrouter.ai/api/v1/chat/completions",
        headers={
            "Authorization": f"Bearer {settings.openrouter_api_key}",
            "Content-Type": "application/json",
            "HTTP-Referer": "https://sgc-legal-ai.vercel.app",
            "X-Title": "SGC Legal AI"
        },
        json={
            "model": settings.model_file_processor,
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": "Распознай и извлеки весь текст с этого изображения. Выведи только распознанный текст, сохраняя структуру и форматирование. Если текста нет, напиши 'Текст не обнаружен'."
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/png;base64,{image_base64}"
                            }
                        }
                    ]
                }
            ],
            "max_tokens": 4096
        },
        timeout=120
    )

    response.raise_for_status()
    result = response.json()
    return result["choices"][0]["message"]["content"]


async def extract_image_gemini(content: bytes, filename: str) -> str:
    """
    Извлечь текст из изображения через Gemini (OpenRouter)
    Gemini поддерживает изображения как multimodal input
    """
    # Кодируем изображение в base64
    image_base64 = base64.b64encode(content).decode('utf-8')
    mime_type = get_image_mime_type(filename)

    # Формируем запрос к Gemini через OpenRouter
    response = requests.post(
        "https://openrouter.ai/api/v1/chat/completions",
        headers={
            "Authorization": f"Bearer {settings.openrouter_api_key}",
            "Content-Type": "application/json",
            "HTTP-Referer": "https://sgc-legal-ai.vercel.app",
            "X-Title": "SGC Legal AI"
        },
        json={
            "model": settings.model_file_processor,
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": "Распознай и извлеки весь текст с этого изображения. Выведи только распознанный текст, сохраняя структуру и форматирование. Если текста нет, напиши 'Текст не обнаружен'."
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:{mime_type};base64,{image_base64}"
                            }
                        }
                    ]
                }
            ],
            "max_tokens": 4096
        },
        timeout=120
    )

    response.raise_for_status()
    result = response.json()

    return result["choices"][0]["message"]["content"]


async def transcribe_audio_gemini(content: bytes, filename: str) -> str:
    """
    Транскрибировать аудио через Gemini (OpenRouter)
    Gemini поддерживает аудио как multimodal input
    """
    # Кодируем аудио в base64
    audio_base64 = base64.b64encode(content).decode('utf-8')
    mime_type = get_audio_mime_type(filename)

    # Формируем запрос к Gemini через OpenRouter
    response = requests.post(
        "https://openrouter.ai/api/v1/chat/completions",
        headers={
            "Authorization": f"Bearer {settings.openrouter_api_key}",
            "Content-Type": "application/json",
            "HTTP-Referer": "https://sgc-legal-ai.vercel.app",
            "X-Title": "SGC Legal AI"
        },
        json={
            "model": settings.model_file_processor,
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": "Транскрибируй это аудио на русском языке. Выведи только текст транскрипции, без комментариев."
                        },
                        {
                            "type": "audio_url",
                            "audio_url": {
                                "url": f"data:{mime_type};base64,{audio_base64}"
                            }
                        }
                    ]
                }
            ],
            "max_tokens": 4096
        },
        timeout=120
    )

    response.raise_for_status()
    result = response.json()

    return result["choices"][0]["message"]["content"]


def get_file_summary(text: str, file_type: str, filename: str) -> str:
    """Создать краткое описание загруженного файла"""
    word_count = len(text.split())
    char_count = len(text)

    type_names = {
        'document': 'документ',
        'pdf': 'PDF-документ',
        'spreadsheet': 'таблица Excel',
        'text': 'текстовый файл',
        'image': 'изображение (OCR)',
        'audio': 'аудиозапись (транскрипция)'
    }

    type_name = type_names.get(file_type, 'файл')

    return f"Загружен {type_name}: {filename} | {word_count} слов, {char_count} символов"
